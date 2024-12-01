# -*- coding: utf-8 -*-


#pip install requests python-docx

import requests
from docx import Document
import os

subscription_key = ""
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus2"
target_language = "pt-br"

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json'
    }

    body = [{
        'text': text
    }]
    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }

    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response_json = response.json()
    return response_json[0]['translations'][0]['text']

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, target_language)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)
  path_translated = path.replace('.docx', '_translated.docx')
  translated_doc.save(path_translated)
  return path_translated

input_file = '/content/letra_talk_to_moon.docx'
translate_document(input_file)

"""#Nova parte"""

#!pip install requests beautifulsoup4 openai langchain-openai

import requests
from bs4 import BeautifulSoup

def extract_text_from_url(url):

  response = requests.get(url)

  if response.status_code == 200:
    soup = BeautifulSoup(response.text, 'html.parser')
    for script_or_style in soup(['script', 'style']):
      script_or_style.decompose()

    texto = soup.get_text(separator= ' ')
    # Limpar texto
    linhas = ({line.strip() for line in texto.splitlines()})
    parts = (phrase.strip() for line in linhas for phrase in line.split(' '))
    texto = '\n'.join(part for part in parts if part)
    return texto
  else:
      print(f"Failed to fetch the URL. Status code: {response.status_code}")
      return Nome
  text = soup.get_text()
  return text


extract_text_from_url('https://dev.to/perssondennis/from-programmer-to-software-developer-the-skills-that-make-the-difference-4k86')

from langchain_openai.chat_models.azure import AzureChatOpenAI

client = AzureChatOpenAI(

    azure_endpoint= "",
    api_key= '',
    api_version= '2024-07-18',
    deployment_name= 'gpt-4o-mini (version:2024-07-18)',
    max_retries= 0
)

def translate_article(text, lang):
  messages = [
      ("system", "Você atua como tradutor de textos"),
      ("user", f'Traduza o {text} para o idioma {lang} e responda em markdown')
      ]

  response = client.invoke(messages)

  print(response.content)
  return response.content

url = "https://dev.to/perssondennis/from-programmer-to-software-developer-the-skills-that-make-the-difference-4k86"
text = extract_text_from_url(url)
translate_article(text, 'pt-br')